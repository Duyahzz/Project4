import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image, ImageDraw, ImageFont
import subprocess
import json
import os

art_dir = "C:/Users/ACER/.gemini/antigravity/brain/dea5ba61-4960-4852-a9d6-f0399d4e1e13"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cover_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def run_query(sql):
    cmd = [
        "sqlcmd", "-S", "localhost,1433", "-U", "sa", "-P", "1",
        "-d", "eStudentDB", "-Q", sql, "-y", "0", "-w", "8000"
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")
    return result.stdout

def clean_json(text):
    start = text.find("[")
    if start != -1:
        end = text.rfind("]")
        if end != -1:
            return text[start:end+1]
    return text

def get_font(size, bold=False):
    try:
        font_name = "arialbd.ttf" if bold else "arial.ttf"
        return ImageFont.truetype(font_name, size)
    except:
        return ImageFont.load_default()

def draw_dfd():
    img = Image.new('RGB', (800, 450), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.ellipse([300, 150, 500, 300], fill=(230, 242, 255), outline=(10, 37, 64), width=3)
    draw.text((360, 210), "eStudiez\nSystem", fill=(10, 37, 64), font=get_font(14, True))
    actors = [
        ("Admin", (50, 180, 170, 260)),
        ("Teacher", (320, 20, 480, 100)),
        ("Student", (630, 180, 750, 260)),
        ("Parent", (320, 350, 480, 430))
    ]
    for label, box in actors:
        draw.rectangle(box, fill=(245, 247, 250), outline=(10, 37, 64), width=2)
        draw.text((box[0]+25, box[1]+25), label, fill=(10, 37, 64), font=get_font(14, True))
    draw.line([170, 220, 300, 220], fill=(10, 37, 64), width=2)
    draw.line([400, 100, 400, 150], fill=(10, 37, 64), width=2)
    draw.line([500, 220, 630, 220], fill=(10, 37, 64), width=2)
    draw.line([400, 300, 400, 350], fill=(10, 37, 64), width=2)
    img.save("scratch/dfd_level0.png")

def draw_architecture():
    # Canvas Size
    width = 800
    height = 600
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Draw light grey grid background
    grid_spacing = 20
    for x in range(0, width, grid_spacing):
        draw.line([x, 0, x, height], fill=(240, 240, 240), width=1)
    for y in range(0, height, grid_spacing):
        draw.line([0, y, width, y], fill=(240, 240, 240), width=1)
        
    # Helpers for rounded boxes
    def draw_rounded_box(draw, box, radius, fill_color, outline_color, width=1):
        x0, y0, x1, y1 = box
        draw.rounded_rectangle([x0, y0, x1, y1], radius=radius, fill=fill_color, outline=outline_color, width=width)

    # ─── ROW 1: PRESENTATION (Thymeleaf & Flutter) ───
    # Icon: Client device (Monitor & Phone)
    draw.polygon([(60, 110), (100, 110), (90, 95), (70, 95)], fill=(120, 144, 156))
    draw.rectangle([75, 95, 85, 105], fill=(144, 164, 174))
    draw_rounded_box(draw, [40, 50, 120, 95], 5, (224, 242, 241), (55, 71, 79), 2)
    draw.text((68, 65), "</>", fill=(38, 50, 56), font=get_font(12, True))
    draw_rounded_box(draw, [110, 40, 140, 95], 4, (230, 242, 255), (0, 112, 192), 2)
    draw.rectangle([112, 45, 138, 85], fill=(255, 255, 255))
    draw.ellipse([123, 88, 127, 92], fill=(0, 112, 192))
    
    # Right box
    draw_rounded_box(draw, [180, 20, 780, 140], 12, (255, 255, 255), (180, 180, 180), 2)
    
    # Thymeleaf Logo & Text
    draw.polygon([(220, 80), (250, 45), (270, 60), (265, 85), (245, 105), (225, 105)], fill=(0, 95, 47))
    draw.polygon([(225, 75), (250, 45), (258, 60), (252, 85), (240, 98)], fill=(0, 120, 60))
    draw.text((280, 50), "Thymeleaf", fill=(50, 50, 50), font=get_font(28, False))
    
    # Flutter Logo & Text
    draw.polygon([(620, 45), (650, 45), (630, 65), (600, 65)], fill=(0, 180, 255))
    draw.polygon([(630, 65), (660, 65), (620, 105), (590, 105)], fill=(0, 112, 192))
    draw.polygon([(620, 105), (660, 65), (645, 65), (605, 105)], fill=(0, 80, 180))
    draw.text((670, 52), "Flutter", fill=(50, 50, 50), font=get_font(28, False))
    
    # ─── ROW 2: SECURITY (JWT) ───
    # Icon: Shield with lock
    draw.polygon([(50, 190), (110, 190), (110, 230), (80, 260), (50, 230)], fill=(245, 247, 250))
    draw.line([(50, 190), (110, 190), (110, 230), (80, 260), (50, 230), (50, 190)], fill=(0, 0, 0), width=3)
    draw_rounded_box(draw, [70, 215, 90, 240], 2, (0, 0, 0), (0, 0, 0), 1)
    draw.arc([73, 205, 87, 220], 180, 360, fill=(0, 0, 0), width=2)
    
    # Right box
    draw_rounded_box(draw, [180, 160, 780, 280], 12, (255, 255, 255), (180, 180, 180), 2)
    
    # JWT Logo & Text
    star_x, star_y = 360, 220
    colors = [
        (251, 0, 137), (218, 0, 137), (137, 0, 218), (0, 180, 255),
        (0, 210, 180), (0, 210, 80), (180, 210, 0), (251, 140, 0)
    ]
    import math
    for i in range(12):
        angle = i * (360 / 12)
        rad = math.radians(angle)
        x_start = star_x + int(10 * math.cos(rad))
        y_start = star_y + int(10 * math.sin(rad))
        x_end = star_x + int(32 * math.cos(rad))
        y_end = star_y + int(32 * math.sin(rad))
        draw.line([x_start, y_start, x_end, y_end], fill=colors[i % len(colors)], width=6)
    draw.line([star_x, star_y - 35, star_x, star_y + 35], fill=(255, 255, 255), width=6)
    draw.text((430, 190), "JWT", fill=(0, 0, 0), font=get_font(42, True))
    
    # ─── ROW 3: BACKEND (Spring Boot) ───
    # Icon: Browser with gear
    draw.rectangle([45, 340, 115, 395], outline=(0, 112, 192), width=2)
    draw.line([45, 350, 115, 350], fill=(0, 112, 192), width=2)
    draw.ellipse([85, 375, 110, 400], outline=(55, 71, 79), width=2)
    for i in range(8):
        angle = i * (360 / 8)
        rad = math.radians(angle)
        gx = 97 + int(15 * math.cos(rad))
        gy = 387 + int(15 * math.sin(rad))
        draw.line([97, 387, gx, gy], fill=(55, 71, 79), width=3)
    draw.ellipse([92, 382, 102, 392], fill=(255, 255, 255), outline=(55, 71, 79), width=2)
    draw.text((55, 402), "BACKEND", fill=(0, 112, 192), font=get_font(8, True))
    
    # Right box
    draw_rounded_box(draw, [180, 300, 780, 420], 12, (255, 255, 255), (180, 180, 180), 2)
    
    # Green Spring Boot Block
    draw.rectangle([400, 301, 620, 419], fill=(109, 179, 63))
    draw.ellipse([420, 335, 470, 385], fill=(255, 255, 255))
    draw.polygon([(435, 368), (455, 342), (460, 355), (452, 373)], fill=(109, 179, 63))
    draw.polygon([(435, 368), (446, 355), (452, 358), (442, 372)], fill=(80, 140, 40))
    draw.text((485, 330), "spring", fill=(255, 255, 255), font=get_font(22, False))
    draw.text((485, 358), "Boot", fill=(255, 255, 255), font=get_font(24, True))
    
    # ─── ROW 4: DATABASE (SQL Server & Firebase) ───
    # Icon: Database Cylinder
    db_x = 80
    draw.ellipse([db_x - 35, 470, db_x + 35, 490], fill=(245, 247, 250), outline=(0, 0, 0), width=3)
    draw.rectangle([db_x - 35, 480, db_x + 35, 510], fill=(245, 247, 250))
    draw.line([db_x - 35, 480, db_x - 35, 510], fill=(0, 0, 0), width=3)
    draw.line([db_x + 35, 480, db_x + 35, 510], fill=(0, 0, 0), width=3)
    draw.ellipse([db_x - 35, 500, db_x + 35, 520], fill=(245, 247, 250), outline=(0, 0, 0), width=3)
    draw.rectangle([db_x - 35, 510, db_x + 35, 540], fill=(245, 247, 250))
    draw.line([db_x - 35, 510, db_x - 35, 540], fill=(0, 0, 0), width=3)
    draw.line([db_x + 35, 510, db_x + 35, 540], fill=(0, 0, 0), width=3)
    draw.ellipse([db_x - 35, 530, db_x + 35, 550], fill=(245, 247, 250), outline=(0, 0, 0), width=3)
    
    # Right box
    draw_rounded_box(draw, [180, 440, 780, 560], 12, (255, 255, 255), (180, 180, 180), 2)
    
    # SQL Server Logo & Text
    px, py = 320, 485
    draw.polygon([(px, py), (px - 25, py + 35), (px + 25, py + 35)], outline=(218, 37, 29), width=2)
    draw.line([px, py, px, py + 35], fill=(218, 37, 29), width=2)
    draw.line([px - 25, py + 35, px + 25, py + 35], fill=(218, 37, 29), width=2)
    draw.line([px, py, px - 12, py + 35], fill=(218, 37, 29), width=1)
    draw.line([px, py, px + 12, py + 35], fill=(218, 37, 29), width=1)
    draw.line([px - 12, py + 18, px + 12, py + 18], fill=(218, 37, 29), width=1)
    draw.text((270, 522), "Microsoft", fill=(80, 80, 80), font=get_font(10, False))
    draw.text((270, 532), "SQL Server", fill=(0, 0, 0), font=get_font(18, True))
    
    # Firebase Blue Block
    draw.rectangle([570, 441, 690, 559], fill=(2, 136, 209))
    fx, fy = 630, 490
    draw.polygon([(fx, fy - 25), (fx - 15, fy + 15), (fx + 15, fy + 15)], fill=(255, 179, 0))
    draw.polygon([(fx, fy - 25), (fx - 5, fy + 15), (fx + 15, fy + 15)], fill=(255, 202, 40))
    draw.text((605, 515), "Firebase", fill=(255, 255, 255), font=get_font(12, True))
    
    img.save("scratch/architecture.png")

def draw_usecase():
    img = Image.new('RGB', (800, 600), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([200, 30, 600, 570], outline=(10, 37, 64), width=2)
    draw.text((320, 40), "System Boundary: eStudiez Portal", fill=(10, 37, 64), font=get_font(13, True))

    def draw_stick(name, x, y):
        draw.ellipse([x-12, y-35, x+12, y-11], fill=(245,247,250), outline=(10,37,64), width=2)
        draw.line([x, y-11, x, y+20], fill=(10,37,64), width=2)
        draw.line([x-20, y-2, x+20, y-2], fill=(10,37,64), width=2)
        draw.line([x, y+20, x-12, y+45], fill=(10,37,64), width=2)
        draw.line([x, y+20, x+12, y+45], fill=(10,37,64), width=2)
        draw.text((x-25, y+55), name, fill=(10, 37, 64), font=get_font(12, True))

    def draw_uc(text, x, y):
        draw.ellipse([x-90, y-18, x+90, y-18+36], fill=(230, 242, 255), outline=(30, 58, 138), width=2)
        draw.text((x-70, y-7), text, fill=(10, 37, 64), font=get_font(11, False))

    draw_stick("Admin", 80, 150)
    draw_stick("Teacher", 80, 420)
    draw_stick("Student", 720, 150)
    draw_stick("Parent", 720, 420)

    ucs = [
        ("Login / Logout", 400, 80),
        ("View Timetable", 400, 150),
        ("View Marks & Att", 400, 220),
        ("Record Attendance", 400, 290),
        ("Manage Marks", 400, 360),
        ("AI Evaluation", 400, 430),
        ("Manage Users", 400, 500),
        ("School News & Chat", 400, 550)
    ]
    for text, x, y in ucs:
        draw_uc(text, x, y)

    # Connections
    for y in [80, 500, 550]: draw.line([80, 150, 310, y], fill=(150, 150, 150), width=1)
    for y in [80, 290, 360, 430, 550]: draw.line([80, 420, 310, y], fill=(150, 150, 150), width=1)
    for y in [80, 150, 220, 550]: draw.line([720, 150, 490, y], fill=(150, 150, 150), width=1)
    for y in [80, 150, 220, 550]: draw.line([720, 420, 490, y], fill=(150, 150, 150), width=1)

    img.save("scratch/usecase_diagram.png")

def draw_generic_sequence(filename, title, lifelines, messages):
    img = Image.new('RGB', (700, 450), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((20, 10), title, fill=(10, 37, 64), font=get_font(16, True))
    
    x_coords = {}
    num_lifelines = len(lifelines)
    step = 600 / (num_lifelines - 1) if num_lifelines > 1 else 300
    for idx, name in enumerate(lifelines):
        x = 50 + idx * step
        x_coords[name] = x
        draw.rectangle((x-45, 40, x+45, 70), fill=(245, 247, 250), outline=(10, 37, 64))
        draw.text((x-30, 48), name, fill=(10, 37, 64), font=get_font(11, True))
        draw.line([x, 70, x, 420], fill=(10, 37, 64), width=1)
        
    y = 110
    for from_actor, to_actor, msg_text in messages:
        x_from = x_coords[from_actor]
        x_to = x_coords[to_actor]
        draw.line([x_from, y, x_to, y], fill=(10, 37, 64), width=2)
        direction = 1 if x_to > x_from else -1
        draw.polygon([x_to, y, x_to - direction*8, y-5, x_to - direction*8, y+5], fill=(10, 37, 64))
        draw.text((min(x_from, x_to) + 15, y - 18), msg_text, fill=(10, 37, 64), font=get_font(10, False))
        y += 42
        
    img.save(filename)

def draw_erd():
    img = Image.new('RGB', (850, 550), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    def draw_entity(name, x, y, w=120, h=40):
        draw.rectangle([x, y, x+w, y+h], fill=(245, 247, 250), outline=(10, 37, 64), width=2)
        draw.text((x+15, y+15), name, fill=(10, 37, 64), font=get_font(12, True))

    draw_entity("Users", 360, 40)
    draw_entity("Students", 150, 140)
    draw_entity("Teachers", 360, 140)
    draw_entity("Parents", 570, 140)
    draw_entity("Classes", 150, 260)
    draw_entity("TimetableSlots", 360, 260, w=130)
    draw_entity("Assessments", 570, 260)
    draw_entity("LessonSessions", 360, 380, w=130)
    draw_entity("AttendanceRecords", 120, 380, w=160)
    draw_entity("StudentMarks", 570, 380)

    draw.line([420, 80, 420, 140], fill=(10, 37, 64), width=1)
    draw.line([420, 100, 210, 100], fill=(10, 37, 64), width=1)
    draw.line([210, 100, 210, 140], fill=(10, 37, 64), width=1)
    draw.line([420, 100, 630, 100], fill=(10, 37, 64), width=1)
    draw.line([630, 100, 630, 140], fill=(10, 37, 64), width=1)
    draw.line([210, 180, 210, 260], fill=(10, 37, 64), width=1)
    draw.line([270, 280, 360, 280], fill=(10, 37, 64), width=1)
    draw.line([420, 180, 420, 260], fill=(10, 37, 64), width=1)
    draw.line([425, 300, 425, 380], fill=(10, 37, 64), width=1)
    draw.line([360, 400, 280, 400], fill=(10, 37, 64), width=1)
    draw.line([170, 180, 170, 380], fill=(10, 37, 64), width=1)
    draw.line([630, 180, 630, 260], fill=(10, 37, 64), width=1)
    draw.line([630, 300, 630, 380], fill=(10, 37, 64), width=1)
    draw.line([150, 160, 100, 160], fill=(10, 37, 64), width=1)
    draw.line([100, 160, 100, 480], fill=(10, 37, 64), width=1)
    draw.line([100, 480, 630, 480], fill=(10, 37, 64), width=1)
    draw.line([630, 480, 630, 420], fill=(10, 37, 64), width=1)

    img.save("scratch/erd_diagram.png")

def draw_database_diagram():
    img = Image.new('RGB', (1000, 750), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    def draw_table_schema(name, cols, x, y, w=180, h=130):
        draw.rectangle([x, y, x+w, y+22], fill=(30, 58, 138), outline=(10, 37, 64), width=1)
        draw.text((x+10, y+5), name, fill=(255, 255, 255), font=get_font(11, True))
        draw.rectangle([x, y+22, x+w, y+h], fill=(250, 250, 250), outline=(10, 37, 64), width=1)
        for idx, col in enumerate(cols):
            draw.text((x+10, y+28 + idx*15), col, fill=(51, 51, 51), font=get_font(9, False))

    draw_table_schema("Users", ["UserId (PK, uniqueid)", "RoleId (FK, int)", "Username (nvarchar)", "PasswordHash (nvarchar)", "FullName (nvarchar)"], 380, 30, h=110)
    draw_table_schema("Students", ["StudentId (PK, uniqueid)", "UserId (FK, uniqueid)", "StudentCode (nvarchar)", "Status (nvarchar)"], 80, 180, h=95)
    draw_table_schema("Teachers", ["TeacherId (PK, uniqueid)", "UserId (FK, uniqueid)", "EmployeeCode (nvarchar)", "SubjectId (FK, int)"], 380, 180, h=95)
    draw_table_schema("Parents", ["ParentId (PK, uniqueid)", "UserId (FK, uniqueid)", "Occupation (nvarchar)"], 680, 180, h=80)
    draw_table_schema("Classes", ["ClassId (PK, int)", "SchoolYearId (int)", "GradeId (int)", "Name (nvarchar)", "HomeroomTeacherId (FK)"], 80, 320, h=110)
    draw_table_schema("TimetableSlots", ["TimetableSlotId (PK, int)", "ClassId (FK, int)", "SubjectId (FK, int)", "TeacherId (FK, uniqueid)", "SemesterId (int)"], 380, 320, h=110)
    draw_table_schema("Assessments", ["AssessmentId (PK, int)", "ClassId (FK, int)", "SubjectId (FK, int)", "TeacherId (FK)", "SemesterId (int)"], 680, 320, h=110)
    draw_table_schema("LessonSessions", ["LessonSessionId (PK, int)", "TimetableSlotId (FK)", "ClassId (FK)", "SubjectId (FK)", "SessionDate (date)"], 380, 480, h=110)
    draw_table_schema("AttendanceRecords", ["AttendanceId (PK, int)", "LessonSessionId (FK)", "StudentId (FK)", "Status (nvarchar)"], 80, 480, h=95)
    draw_table_schema("StudentMarks", ["StudentMarkId (PK, int)", "AssessmentId (FK)", "StudentId (FK)", "Score (decimal)"], 680, 480, h=95)

    # Connections
    draw.line([470, 140, 470, 180], fill=(10, 37, 64), width=1)
    draw.line([470, 155, 170, 155], fill=(10, 37, 64), width=1)
    draw.line([170, 155, 170, 180], fill=(10, 37, 64), width=1)
    draw.line([470, 155, 770, 155], fill=(10, 37, 64), width=1)
    draw.line([770, 155, 770, 180], fill=(10, 37, 64), width=1)
    draw.line([170, 275, 170, 480], fill=(10, 37, 64), width=1)
    draw.line([260, 350, 380, 350], fill=(10, 37, 64), width=1)
    draw.line([470, 275, 470, 320], fill=(10, 37, 64), width=1)
    draw.line([470, 430, 470, 480], fill=(10, 37, 64), width=1)
    draw.line([380, 530, 260, 530], fill=(10, 37, 64), width=1)
    draw.line([770, 430, 770, 480], fill=(10, 37, 64), width=1)

    img.save("scratch/database_diagram.png")

def get_db_schema():
    cols_sql = """
    SELECT 
        c.TABLE_NAME, 
        c.COLUMN_NAME, 
        c.DATA_TYPE, 
        COALESCE(CAST(c.CHARACTER_MAXIMUM_LENGTH AS VARCHAR(10)), '') AS CHARACTER_MAXIMUM_LENGTH, 
        c.IS_NULLABLE, 
        COALESCE(c.COLUMN_DEFAULT, '') AS COLUMN_DEFAULT
    FROM INFORMATION_SCHEMA.COLUMNS c
    FOR JSON PATH;
    """
    pks_sql = """
    SELECT 
        ku.TABLE_NAME,
        ku.COLUMN_NAME
    FROM INFORMATION_SCHEMA.KEY_COLUMN_USAGE ku
    INNER JOIN INFORMATION_SCHEMA.TABLE_CONSTRAINTS tc ON ku.CONSTRAINT_NAME = tc.CONSTRAINT_NAME
    WHERE tc.CONSTRAINT_TYPE = 'PRIMARY KEY'
    FOR JSON PATH;
    """
    fks_sql = "SELECT tp.name AS TABLE_NAME, cp.name AS COLUMN_NAME, tr.name AS REFERENCED_TABLE_NAME FROM sys.foreign_keys fk INNER JOIN sys.foreign_key_columns fkc ON fk.object_id = fkc.constraint_object_id INNER JOIN sys.tables tp ON fkc.parent_object_id = tp.object_id INNER JOIN sys.columns cp ON fkc.parent_object_id = cp.object_id AND fkc.parent_column_id = cp.column_id INNER JOIN sys.tables tr ON fkc.referenced_object_id = tr.object_id FOR JSON PATH;"
    try:
        cols_raw = run_query(cols_sql)
        pks_raw = run_query(pks_sql)
        fks_raw = run_query(fks_sql)
        print("Cols Raw Len:", len(cols_raw))
        print("Pks Raw Len:", len(pks_raw))
        print("Fks Raw Len:", len(fks_raw))
        print("Fks Raw Content:", repr(fks_raw[:500]))
        cols_out = json.loads(clean_json(cols_raw), strict=False)
        pks_out = json.loads(clean_json(pks_raw), strict=False)
        fks_out = json.loads(clean_json(fks_raw), strict=False)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print("Fallback schema used due to DB error:", e)
        return {}, {}, {}
    pks = {}
    for item in pks_out:
        t = item.get("TABLE_NAME")
        c = item.get("COLUMN_NAME")
        if t not in pks: pks[t] = set()
        pks[t].add(c)
    fks = {}
    for item in fks_out:
        t = item.get("TABLE_NAME")
        c = item.get("COLUMN_NAME")
        ref = item.get("REFERENCED_TABLE_NAME")
        if t not in fks: fks[t] = {}
        fks[t][c] = ref
    tables = {}
    for item in cols_out:
        t = item.get("TABLE_NAME")
        if t not in tables: tables[t] = []
        tables[t].append(item)
    return tables, pks, fks

def main():
    print("Drawing diagram images programmatically...")
    draw_dfd()
    draw_architecture()
    draw_usecase()
    draw_erd()
    draw_database_diagram()

    # Draw all 12 sequence diagrams
    draw_generic_sequence("scratch/seq_1.png", "Sequence: User Login", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "Login Request"), ("Controller", "OAuth2/Auth", "Create Token"),
                           ("OAuth2/Auth", "MonolithService", "Call Login & Check Role"), ("MonolithService", "Database", "Find User"),
                           ("Database", "MonolithService", "Found User"), ("MonolithService", "OAuth2/Auth", "Return Account"),
                           ("OAuth2/Auth", "Controller", "Return Token"), ("Controller", "User", "Redirect Dashboard")])
                           
    draw_generic_sequence("scratch/seq_2.png", "Sequence: User Logout", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService"],
                          [("User", "Controller", "Log Out Request"), ("Controller", "OAuth2/Auth", "Check Token"),
                           ("OAuth2/Auth", "MonolithService", "Call Logout Function"), ("MonolithService", "OAuth2/Auth", "Logout Success"),
                           ("OAuth2/Auth", "Controller", "Cancel Token"), ("Controller", "User", "Redirect Login")])

    draw_generic_sequence("scratch/seq_3.png", "Sequence: Change Password", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "Change Password Request"), ("Controller", "OAuth2/Auth", "Check Token"),
                           ("OAuth2/Auth", "MonolithService", "Call Change Pwd Function"), ("MonolithService", "Database", "Update Pwd Hash"),
                           ("Database", "MonolithService", "Pwd Updated Success"), ("MonolithService", "Controller", "Return Success"),
                           ("Controller", "User", "Display Success Message")])

    draw_generic_sequence("scratch/seq_4.png", "Sequence: Sign Up", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "Register Account"), ("Controller", "OAuth2/Auth", "Create Token"),
                           ("OAuth2/Auth", "MonolithService", "Call Sign Up"), ("MonolithService", "Database", "Create User Row"),
                           ("Database", "MonolithService", "Row Created Success"), ("MonolithService", "OAuth2/Auth", "Welcome Email"),
                           ("OAuth2/Auth", "Controller", "Success"), ("Controller", "User", "Show Login Screen")])

    draw_generic_sequence("scratch/seq_5.png", "Sequence: View Profile", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "View Profile Click"), ("Controller", "OAuth2/Auth", "Check Token"),
                           ("OAuth2/Auth", "MonolithService", "Get Profile details"), ("MonolithService", "Database", "Select User details"),
                           ("Database", "MonolithService", "Return row data"), ("MonolithService", "Controller", "Return Profile JSON"),
                           ("Controller", "User", "Display profile details")])

    draw_generic_sequence("scratch/seq_6.png", "Sequence: Edit Profile", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "Edit Profile details"), ("Controller", "OAuth2/Auth", "Check Token"),
                           ("OAuth2/Auth", "MonolithService", "Update Profile details"), ("MonolithService", "Database", "Update user set info"),
                           ("Database", "MonolithService", "Return success"), ("MonolithService", "Controller", "Return updated profile"),
                           ("Controller", "User", "Display saved profile")])

    draw_generic_sequence("scratch/seq_7.png", "Sequence: View / Search Student", 
                          ["Admin", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("Admin", "Controller", "Search Student request"), ("Controller", "OAuth2/Auth", "Check Admin Token"),
                           ("OAuth2/Auth", "MonolithService", "Search Student list"), ("MonolithService", "Database", "Select Students"),
                           ("Database", "MonolithService", "Return records"), ("MonolithService", "Controller", "Return list JSON"),
                           ("Controller", "Admin", "Display search list")])

    draw_generic_sequence("scratch/seq_8.png", "Sequence: View / Search Teacher", 
                          ["Admin", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("Admin", "Controller", "Search Teacher request"), ("Controller", "OAuth2/Auth", "Check Admin Token"),
                           ("OAuth2/Auth", "MonolithService", "Search Teacher list"), ("MonolithService", "Database", "Select Teachers"),
                           ("Database", "MonolithService", "Return records"), ("MonolithService", "Controller", "Return list JSON"),
                           ("Controller", "Admin", "Display search list")])

    draw_generic_sequence("scratch/seq_9.png", "Sequence: Add Class", 
                          ["Admin", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("Admin", "Controller", "Add Class Request"), ("Controller", "OAuth2/Auth", "Check Admin Token"),
                           ("OAuth2/Auth", "MonolithService", "Create Class record"), ("MonolithService", "Database", "Insert Class row"),
                           ("Database", "MonolithService", "Success"), ("MonolithService", "Controller", "Return Class details"),
                           ("Controller", "Admin", "Show Class added toast")])

    draw_generic_sequence("scratch/seq_10.png", "Sequence: Edit Class", 
                          ["Admin", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("Admin", "Controller", "Edit Class Request"), ("Controller", "OAuth2/Auth", "Check Admin Token"),
                           ("OAuth2/Auth", "MonolithService", "Update Class record"), ("MonolithService", "Database", "Update Class row"),
                           ("Database", "MonolithService", "Success"), ("MonolithService", "Controller", "Return Success"),
                           ("Controller", "Admin", "Show Class edited toast")])

    draw_generic_sequence("scratch/seq_11.png", "Sequence: View Timetable / Marks", 
                          ["User", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("User", "Controller", "View Timetable/Marks grid"), ("Controller", "OAuth2/Auth", "Check Token"),
                           ("OAuth2/Auth", "MonolithService", "Fetch Timetable/Marks"), ("MonolithService", "Database", "Select records"),
                           ("Database", "MonolithService", "Return rows list"), ("MonolithService", "Controller", "Return records JSON"),
                           ("Controller", "User", "Display Timetable/Marks UI")])

    draw_generic_sequence("scratch/seq_12.png", "Sequence: Record Attendance / Marks", 
                          ["Teacher", "Controller", "OAuth2/Auth", "MonolithService", "Database"],
                          [("Teacher", "Controller", "Record Attendance/Marks"), ("Controller", "OAuth2/Auth", "Check Teacher Token"),
                           ("OAuth2/Auth", "MonolithService", "Save Attendance/Marks details"), ("MonolithService", "Database", "Insert/Update records"),
                           ("Database", "MonolithService", "Confirm save success"), ("MonolithService", "Controller", "Success"),
                           ("Controller", "Teacher", "Display Success Toast")])

    print("Fetching database catalog...")
    db_tables, pks, fks = get_db_schema()
    if not db_tables:
        print("Error: Could not retrieve database catalog.")
        return

    doc = docx.Document()

    # 1-inch margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    PRIMARY_COLOR = RGBColor(10, 37, 64)
    SECONDARY_COLOR = RGBColor(30, 58, 138)
    DARK_TEXT = RGBColor(51, 51, 51)

    # Helper for headings and text
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = PRIMARY_COLOR
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = SECONDARY_COLOR
        return p

    def add_body(text, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_TEXT
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_TEXT
        return p

    def add_review_header(review_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(150)
        p.paragraph_format.space_after = Pt(150)
        r = p.add_run(review_text)
        r.bold = True
        r.font.size = Pt(36)
        r.font.color.rgb = PRIMARY_COLOR
        doc.add_page_break()

    # ─── COVER PAGE ───
    # ─── GENERATE APTECH LOGO IMAGE ───
    logo_path = "scratch/aptech_logo.png"
    logo_img = Image.new('RGB', (350, 150), 'white')
    logo_draw = ImageDraw.Draw(logo_img)
    
    # Draw black block
    logo_draw.rectangle([10, 10, 240, 75], fill=(26, 26, 26))
    
    # Draw red sloped polygon
    logo_draw.polygon([(10, 10), (65, 10), (45, 75), (10, 75)], fill=(218, 37, 29))
    
    try:
        font_aptech = ImageFont.truetype("arialbd.ttf", 26)
        font_sub = ImageFont.truetype("arial.ttf", 9)
        font_alliance = ImageFont.truetype("arialbd.ttf", 9)
        font_fpt = ImageFont.truetype("arialbd.ttf", 10)
    except Exception:
        font_aptech = ImageFont.load_default()
        font_sub = ImageFont.load_default()
        font_alliance = ImageFont.load_default()
        font_fpt = ImageFont.load_default()
        
    logo_draw.text((70, 20), "Aptech", fill=(255, 255, 255), font=font_aptech)
    
    # Paw mark
    logo_draw.ellipse([185, 20, 192, 27], fill=(247, 148, 29))
    logo_draw.ellipse([194, 17, 202, 25], fill=(247, 148, 29))
    logo_draw.ellipse([204, 20, 211, 27], fill=(247, 148, 29))
    logo_draw.ellipse([188, 30, 208, 44], fill=(247, 148, 29))
    
    # "COMPUTER EDUCATION"
    logo_draw.text((70, 52), "COMPUTER EDUCATION", fill=(247, 148, 29), font=font_sub)
    
    # "Alliance with"
    logo_draw.text((10, 85), "Alliance with", fill=(10, 37, 64), font=font_alliance)
    
    # FPT colored logo
    logo_draw.text((80, 85), "F", fill=(247, 148, 29), font=font_fpt)
    logo_draw.text((90, 85), "P", fill=(0, 176, 80), font=font_fpt)
    logo_draw.text((100, 85), "T", fill=(0, 112, 192), font=font_fpt)
    
    # "Education"
    logo_draw.text((115, 85), "Education", fill=(46, 117, 182), font=font_alliance)
    
    logo_img.save(logo_path)

    # ─── HEADER TABLE ───
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(header_table)
    
    # Set widths for header table columns
    cell_hdr_0 = header_table.cell(0, 0)
    cell_hdr_1 = header_table.cell(0, 1)
    cell_hdr_0.width = Inches(2.2)
    cell_hdr_1.width = Inches(4.3)
    set_cell_margins(cell_hdr_0, top=0, bottom=0, left=0, right=0)
    set_cell_margins(cell_hdr_1, top=0, bottom=0, left=0, right=0)
    
    p_hdr_left = cell_hdr_0.paragraphs[0]
    p_hdr_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_logo = p_hdr_left.add_run()
    r_logo.add_picture(logo_path, width=Inches(1.8))
    
    p_hdr_right = cell_hdr_1.paragraphs[0]
    p_hdr_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_hdr_1 = p_hdr_right.add_run("FPT ACADEMY INTERNATIONAL\n")
    r_hdr_1.bold = True
    r_hdr_1.font.name = "Arial"
    r_hdr_1.font.size = Pt(11)
    r_hdr_1.font.color.rgb = RGBColor(0, 0, 0)
    r_hdr_2 = p_hdr_right.add_run("FPT – APTECH COMPUTER EDUCATION")
    r_hdr_2.bold = True
    r_hdr_2.font.name = "Arial"
    r_hdr_2.font.size = Pt(11)
    r_hdr_2.font.color.rgb = RGBColor(0, 0, 0)
    
    # ─── CENTER DETAILS ───
    p_center = doc.add_paragraph()
    p_center.paragraph_format.space_before = Pt(30)
    p_center.paragraph_format.space_after = Pt(2)
    r_name = p_center.add_run("Center Name: ACE-HCMC-2-FPT.")
    r_name.bold = True
    r_name.font.name = "Arial"
    r_name.font.size = Pt(11.5)
    r_name.font.color.rgb = RGBColor(0, 0, 0)

    p_addr = doc.add_paragraph()
    p_addr.paragraph_format.space_after = Pt(70) # large space before title
    r_addr = p_addr.add_run("Address: 21 Bis Hậu Giang Street, Tân Sơn Nhất Ward, Ho Chi Minh City, Viet Nam.")
    r_addr.bold = True
    r_addr.font.name = "Arial"
    r_addr.font.size = Pt(11.5)
    r_addr.font.color.rgb = RGBColor(0, 0, 0)
    
    # ─── MAIN TITLE ───
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("eStudiez")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(52)
    r_title.font.color.rgb = RGBColor(237, 125, 49) # Orange color

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(60) # space before table
    r_sub = p_sub.add_run("Design Document")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(24)
    r_sub.font.color.rgb = RGBColor(0, 0, 0)
    
    # ─── INFO & MEMBERS TABLE ───
    table_info = doc.add_table(rows=8, cols=3)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cover_table_borders(table_info)

    col_widths = [Inches(1.5), Inches(3.2), Inches(1.8)]

    # Merge Col 1 & Col 2 for the first 4 rows
    for r_idx in range(4):
        c1 = table_info.cell(r_idx, 1)
        c2 = table_info.cell(r_idx, 2)
        c1.merge(c2)

    row_data_merged = [
        ("Supervisor:", "Mr. Trần Phước Sinh"),
        ("Semester:", "4"),
        ("Batch No:", "T1.2406.E1"),
        ("Group No:", "5")
    ]

    for r_idx, (lbl, val) in enumerate(row_data_merged):
        cell_lbl = table_info.cell(r_idx, 0)
        set_cell_margins(cell_lbl, top=120, bottom=120, left=150, right=150)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_lbl = p_lbl.add_run(lbl)
        r_lbl.bold = True
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(10.5)
        r_lbl.font.color.rgb = RGBColor(0, 0, 0)

        cell_val = table_info.cell(r_idx, 1)
        set_cell_margins(cell_val, top=120, bottom=120, left=150, right=150)
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_val = p_val.add_run(val)
        r_val.font.name = "Arial"
        r_val.font.size = Pt(10.5)
        r_val.font.color.rgb = RGBColor(0, 0, 0)

    # Row 4 (Table header)
    headers = [("Order:", WD_ALIGN_PARAGRAPH.LEFT), ("Full name", WD_ALIGN_PARAGRAPH.LEFT), ("Roll No.", WD_ALIGN_PARAGRAPH.LEFT)]
    for c_idx, (text, align) in enumerate(headers):
        cell = table_info.cell(4, c_idx)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Rows 5 to 7 (Members list)
    members_list = [
        ("1.", "Đỗ Nguyễn Thiện Hoàng", "Student1545463"),
        ("2.", "Trần Quang Khải", "Student1457498"),
        ("3.", "Phan Văn Duy", "Student1470114")
    ]

    for i, (order, name, roll) in enumerate(members_list):
        r_idx = 5 + i
        # Col 0 (Order)
        c0 = table_info.cell(r_idx, 0)
        set_cell_margins(c0, top=120, bottom=120, left=150, right=150)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(order)
        r0.font.name = "Arial"
        r0.font.size = Pt(10.5)
        r0.font.color.rgb = RGBColor(0, 0, 0)

        # Col 1 (Full name)
        c1 = table_info.cell(r_idx, 1)
        set_cell_margins(c1, top=120, bottom=120, left=150, right=150)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(name)
        r1.font.name = "Arial"
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(0, 0, 0)

        # Col 2 (Roll No.)
        c2 = table_info.cell(r_idx, 2)
        set_cell_margins(c2, top=120, bottom=120, left=150, right=150)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(roll)
        r2.font.name = "Arial"
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(0, 0, 0)

    # Set widths on all cells
    for row in table_info.rows:
        for c_idx, width in enumerate(col_widths):
            try:
                row.cells[c_idx].width = width
            except Exception:
                pass

    doc.add_page_break()

    # ─── TABLE OF CONTENTS ───
    add_h1("Table of Contents")
    toc_lines = [
        ("Acknowledgements", "4"),
        ("REVIEW 1", "5"),
        ("1. Overview", "6"),
        ("  1.1 Project Information", "6"),
        ("  1.2 Product Background", "6"),
        ("  1.3 Existing Systems", "6"),
        ("  1.4 Business Opportunity", "7"),
        ("  1.5 Project Scope & Limitations", "7"),
        ("2. User Requirements", "8"),
        ("  2.1 Actor Descriptions", "8"),
        ("  2.2 Level 0 DFD (Data Flow Diagram)", "9"),
        ("  2.3 High Level Use Case Map", "9"),
        ("  2.4 Use Case List", "10"),
        ("  2.5 Use Case & Actor Mapping", "10"),
        ("3. Management Plan", "11"),
        ("  3.1 Scope and WBS Estimation", "11"),
        ("  3.2 Management Approach", "12"),
        ("  3.3 Training Plan", "12"),
        ("4. System Requirements", "13"),
        ("  4.1 Hardware Requirement", "13"),
        ("  4.2 Software Requirements", "13"),
        ("  4.3 Technology Summary", "13"),
        ("  Task Sheet Review 1", "14"),
        ("REVIEW 2", "15"),
        ("1. Architecture Design", "16"),
        ("2. Use Case Specifications & Sequence Diagrams", "17"),
        ("  2.1 UC101 - Login", "17"),
        ("  2.2 UC102 - Logout", "19"),
        ("  2.3 UC103 - Change Password", "20"),
        ("  2.4 UC104 - Sign Up", "21"),
        ("  2.5 UC105 - View Profile", "23"),
        ("  2.6 UC106 - Edit Profile", "24"),
        ("  2.7 UC201 - View / Search Student", "26"),
        ("  2.8 UC202 - View / Search Teacher", "27"),
        ("  2.9 UC203 - Add Class", "29"),
        ("  2.10 UC204 - Edit Class", "30"),
        ("  2.11 UC205 - View Timetable / Marks", "32"),
        ("  2.12 UC206 - Record Attendance / Marks", "33"),
        ("3. Table Definitions (Database Schema)", "35"),
        ("4. ERD (Entity Relationship Diagram)", "60"),
        ("  Task Sheet Review 2", "61"),
        ("REVIEW 3", "62"),
        ("1. Database Diagram", "63"),
        ("2. Web GUI Design", "64"),
        ("3. Mobile GUI Design", "78"),
        ("  Task Sheet Review 3", "81"),
        ("  Project Task Sheet", "82")
    ]
    for section_title, page_no in toc_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(section_title)
        if "REVIEW" in section_title or section_title == "Acknowledgements" or section_title == "Table of Contents":
            r_title.bold = True
        r_dots = p.add_run(" " + "." * (80 - len(section_title)) + " ")
        r_page = p.add_run(page_no)
        r_page.bold = True
        
    doc.add_page_break()

    # ─── ACKNOWLEDGEMENTS ───
    add_h1("Acknowledgements")
    add_body("The satisfaction that accompanies the successful completion of any task would be incomplete without the mention of people whose ceaseless cooperation made it possible, whose constant guidance and encouragement crown all efforts with success.")
    add_body("We are immensely grateful to our project guide, Mr. Tran Phuoc Sinh, for his guidance, inspiration, and constructive suggestions that significantly contributed to the preparation and successful completion of our study-progress tracking project. His expertise and support were invaluable throughout the process.")
    add_body("We also extend our heartfelt thanks to our colleagues who contributed their time, effort, and insights, which were crucial to the successful completion of this project. Their collaboration and dedication played a pivotal role in overcoming challenges and achieving our goals.")
    
    doc.add_page_break()

    # ─── REVIEW 1 SECTION ───
    add_review_header("REVIEW 1")

    # ─── SECTION 1: OVERVIEW ───
    add_h1("1. Overview")
    
    add_h2("1.1 Project Information")
    add_bullet("Project name: eStudiez (eStudent) Your Next-Gen All-in-One School Progress Tracker")
    add_bullet("Group name: Group 01")
    add_bullet("Software type: Web Application, Mobile Web Portal")
    
    add_h2("1.2 Product Background")
    add_body("Designing and developing a School Study-Progress Tracking Web Application involves several stages of production. Here's a brief overview of the production background for such a project:")
    add_body("Requirement Gathering: This stage involves identifying the requirements of both the school board (Admins), subject teachers, students, and parents/guardians for the application. It is essential to gather feedback from users to determine the features and functionality that should be included in the application.")
    add_body("Design: Once the requirements are clear, the design phase begins. It includes creating the visual design of the application and defining the user interface and user experience. A prototype may be developed to get feedback from users.")
    add_body("Development: This phase involves the actual development of the application. The development team uses the appropriate technology stack to build the application, and the features identified in the design stage will be implemented.")
    add_body("Testing: Before the application is released, it must be thoroughly tested to ensure that it meets the requirements and functions as expected. Various types of testing, including functional, performance, and security testing, should be conducted.")
    
    add_h2("1.3 Existing Systems")
    add_body("There are several existing systems for school portal and study-progress management that we can consider when developing our own app. Here are a few examples:")
    add_body("vnEdu: Developed by VNPT, vnEdu is one of the most popular school portals in Vietnam. It contains online grade books, attendance tracking, and school communication systems. However, it can feel fragmented and lacks real-time classroom group chat functions and AI evaluation support.")
    add_body("SMAS: A school management system developed by Viettel. It is heavy on administrative records and manager statistics but is less streamlined for interactive student-teacher communication and sharing study resources.")
    add_body("Canvas LMS: A global learning management system. While excellent for university courses, it lacks specific localized school administrative structures (such as Vietnamese grade levels, regular classes, and homeroom parent links).")
    
    add_h2("1.4 Business Opportunity")
    add_body("Portal centralization has modernized school operations. Key opportunities include:")
    add_bullet("Centralized Academic Portal: Providing a single platform for tracking subject marks, attendance, and timetables.")
    add_bullet("Parent Engagement: Connecting parents directly with classrooms via Parent-Teacher chat groups and child progress notifications.")
    add_bullet("Resource Accessibility: Online subject resources (documents, videos, links) ensure students can study remotely 24/7.")
    add_bullet("AI Academic Insights: Generating automated learning recommendations for students based on their exam performance.")
    
    add_h2("1.5 Project Scope & Limitations")
    add_body("Scope: Academic progress tracking (marks & attendance), timetables, news bulletins, notifications, class chat groups, extra revision classes, study resources repository, and parent contact management.")
    add_body("Limitations: No teacher HR management (payroll/recruitment), no admissions processing, no student fee transactions, and restricted to high school class structures (grades 10, 11, and 12).")
    doc.add_page_break()

    # ─── SECTION 2: USER REQUIREMENTS ───
    add_h1("2. User Requirements")
    add_h2("2.1 Actor Descriptions")
    table_act = doc.add_table(rows=5, cols=2)
    table_act.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_act)
    
    headers_act = ["Actor", "Description"]
    for col_idx, text in enumerate(headers_act):
        cell = table_act.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        
    actors = [
        ("Admin (Ban giám hiệu)", "Super Admin of the system. Manages user creation (students, teachers, parents), configures classes, maps teacher-class assignments, and posts school-wide news."),
        ("Teacher (Giáo viên)", "Subject teachers who take student attendance per period, manage and record scores for assignments/exams, upload study documents, and write evaluations."),
        ("Student (Học sinh)", "Enrolled students who view timetables, check personal marks and attendance, view/download class resources, and participate in student-teacher chat rooms."),
        ("Parent (Phụ huynh)", "Guardians linked to students. Monitor child's progress, view timetables, read school bulletins, and participate in Parent-Teacher chat groups.")
    ]
    for row_idx, data in enumerate(actors):
        for col_idx, text in enumerate(data):
            cell = table_act.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT
            
    add_h2("2.2 Level 0 DFD (Data Flow Diagram)")
    add_body("Below is the Level 0 DFD illustrating the flow of data between the system actors and eStudiez:")
    doc.add_picture("scratch/dfd_level0.png", width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.add_run("Figure 2.1: Level 0 DFD Diagram").italic = True
    
    add_h2("2.3 High Level Use Case Map")
    add_body("Use Cases are structured into four modules:")
    add_bullet("Common: UC101 (Login/Logout), UC102 (Change Password), UC103 (View/Edit Profile), UC104 (Notifications).")
    add_bullet("Student/Parent: UC201 (View Timetable), UC202 (View Marks & Attendance), UC203 (Download Resources), UC204 (Submit Feedback).")
    add_bullet("Teacher: UC301 (Record Attendance), UC302 (Manage Marks), UC303 (Upload Resources), UC304 (Submit AI Evaluation).")
    add_bullet("Admin: UC401 (Manage Users), UC402 (Configure Classes), UC403 (Assign Teachers), UC404 (Publish News).")
    
    add_body("Below is the high-level UML Use Case diagram mapping actors to their respective functions:")
    doc.add_picture("scratch/usecase_diagram.png", width=Inches(6.0))
    p_cap_uc = doc.add_paragraph()
    p_cap_uc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_uc.add_run("Figure 2.2: High Level Use Case Diagram").italic = True
    
    add_h2("2.4 Use Case List")
    table_uc = doc.add_table(rows=11, cols=3)
    table_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_uc)
    
    headers_uc = ["UC Code", "Use Case Name", "Description"]
    for col_idx, text in enumerate(headers_uc):
        cell = table_uc.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    usecases = [
        ("UC101", "Login / Logout", "Authenticates all users using encrypted passwords."),
        ("UC102", "Change Password", "Allows users to update their credentials for safety."),
        ("UC103", "View/Edit Profile", "View and update email, phone, and profile avatar."),
        ("UC201", "View Timetable", "Students/Parents view class weekly schedule grid."),
        ("UC202", "View Progress", "Students/Parents track marks and attendance status."),
        ("UC203", "Download Resource", "Students browse and download subject study materials."),
        ("UC301", "Record Attendance", "Teachers mark student presence, lateness, or excuses."),
        ("UC302", "Update Marks", "Teachers record exam scores and detailed remarks."),
        ("UC303", "AI Evaluation", "AI processes scores to suggest custom study steps."),
        ("UC401", "Manage Users & Classes", "Admins create students, classes, and assign teachers.")
    ]
    for row_idx, data in enumerate(usecases):
        for col_idx, text in enumerate(data):
            cell = table_uc.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT

    add_h2("2.5 Use Case & Actor mapping")
    table_map = doc.add_table(rows=len(usecases)+1, cols=5)
    table_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_map)
    headers_map = ["Code", "Use Case", "Admin", "Teacher", "Student/Parent"]
    for col_idx, text in enumerate(headers_map):
        cell = table_map.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
    for row_idx, data in enumerate(usecases):
        cell_c = table_map.cell(row_idx + 1, 0)
        cell_n = table_map.cell(row_idx + 1, 1)
        cell_c.paragraphs[0].add_run(data[0])
        cell_n.paragraphs[0].add_run(data[1])
        # Mark mapping Xs
        code = data[0]
        if code in ["UC101", "UC102", "UC103"]:
            for c_idx in [2, 3, 4]: table_map.cell(row_idx + 1, c_idx).paragraphs[0].add_run("X")
        elif code in ["UC201", "UC202", "UC203"]:
            table_map.cell(row_idx + 1, 4).paragraphs[0].add_run("X")
        elif code in ["UC301", "UC302", "UC303"]:
            table_map.cell(row_idx + 1, 3).paragraphs[0].add_run("X")
        elif code in ["UC401"]:
            table_map.cell(row_idx + 1, 2).paragraphs[0].add_run("X")

    add_h1("3. Management Plan")
    add_h2("3.1 Scope and Estimation (WBS)")
    add_body("Project scheduling over 5 weeks (July 10 to August 18). Detailed project tasks breakdown:")
    table_wbs = doc.add_table(rows=10, cols=3)
    table_wbs.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_wbs)
    headers_wbs = ["WBS Task Description", "Complexity", "Est. Effort (days)"]
    for col_idx, text in enumerate(headers_wbs):
        cell = table_wbs.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    wbs_data = [
        ("Phase 1: Project Plan & SRS Documenting", "Medium", "4.0"),
        ("Phase 2: Database Schema & Clustered Indexes Setup", "Medium", "2.0"),
        ("Phase 3: Security Auth filter & JWT tokens", "High", "3.0"),
        ("Phase 4: Subject Marks & Attendance backend controller", "Medium", "3.5"),
        ("Phase 5: Timetable grid & Class resources download API", "Medium", "1.5"),
        ("Phase 6: Frontend context global states setup", "High", "2.0"),
        ("Phase 7: Student Timetable grid & Attendance view panels", "High", "4.0"),
        ("Phase 8: Teacher Grading Accordion & AI auto-fill sheet", "High", "4.5"),
        ("Phase 9: Integration testing, build checks & CI/CD deployment", "Medium", "3.0")
    ]
    for row_idx, row_info in enumerate(wbs_data):
        for col_idx, text in enumerate(row_info):
            cell = table_wbs.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].add_run(text)
            
    add_h2("3.2 Management Approach")
    add_body("The development team uses a DevOps continuous pipeline (Git checkout, Gradle compile, npm build checks) to sync development loops.")
    
    add_h2("3.3 Training Plan")
    add_bullet("Spring Boot & Spring Data JPA: Week 1, 3 days (Mandatory)")
    add_bullet("React & TypeScript context states: Week 2, 3 days (Mandatory)")
    add_bullet("Git branching and GitHub Actions setup: Week 1, 1 day (Mandatory)")
    
    add_h1("4. System Requirements")
    add_bullet("Hardware: Dual-Core 2.0 GHz, 8 GB RAM, 50 GB free SSD storage, Standard internet connection.")
    add_bullet("Software: SQL Server 2019+, IntelliJ IDEA, JDK 17, Node.js LTS v20, Google Chrome browser.")
    add_bullet("Technology Summary: React, TypeScript, Tailwind CSS, Spring Boot REST APIs, JPA Hibernate.")
    
    add_h2("Task Sheet Review 1")
    t_wbs1 = doc.add_table(rows=4, cols=4)
    t_wbs1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_wbs1)
    wbs1_headers = ["Member Name", "Assigned Task", "Period", "Status"]
    for col_idx, text in enumerate(wbs1_headers):
        cell = t_wbs1.cell(0, col_idx)
        set_cell_background(cell, "0A2540")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    wbs1_rows = [
        ("Đỗ Nguyễn Thiện Hoàng", "High-level Use Case, WBS Planning & Timetable Setup", "Week 1", "Completed"),
        ("Trần Quang Khải", "Use Case List, Actor Mapping & Product Overview", "Week 1", "Completed"),
        ("Phan Văn Duy", "System Requirements, Hardware/Software Specs & Scope", "Week 1", "Completed")
    ]
    for row_idx, row_info in enumerate(wbs1_rows):
        for col_idx, text in enumerate(row_info):
            cell = t_wbs1.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].add_run(text)

    doc.add_page_break()

    # ─── REVIEW 2 SECTION ───
    add_review_header("REVIEW 2")

    # ─── SECTION 1: ARCHITECTURE DESIGN ───
    add_h1("1. Architecture Design")
    add_body("Below is the software architecture diagram displaying React client, Spring Boot REST API, and SQL Server DB layers:")
    doc.add_picture("scratch/architecture.png", width=Inches(6.0))
    p_cap_arch = doc.add_paragraph()
    p_cap_arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_arch.add_run("Figure 3.1: eStudiez N-Tier Software Architecture").italic = True
    
    # ─── SECTION 2: USE CASE SPECIFICATIONS & SEQUENCE DIAGRAMS (12 Use Cases) ───
    add_h1("2. Use Case Specifications & Sequence Diagrams")
    
    uc_specs = [
        ("UC101 - Login", "Unauthorized User", "User logs in to access system resources.", 
         "Anonymous User", "Successful login redirects to Dashboard with JWT stored in frontend.",
         "1. User enters registered Username/Email and Password.\n2. Clicks Login.\n3. Server verifies credential hashes.\n4. Generates and returns a JWT token.\n5. Redirects to respective dashboard.", "scratch/seq_1.png"),
         
        ("UC102 - Logout", "Authenticated User", "User leaves the portal session.",
         "User logged in", "Session cleared from browser; redirected to Login page.",
         "1. User clicks the Sign Out button.\n2. Frontend clears the stored JWT from localStorage.\n3. Client state is reset to null.\n4. Screen redirects instantly to the Login form.", "scratch/seq_2.png"),
         
        ("UC103 - Change Password", "Authenticated User", "User retypes password for safety.",
         "User logged in", "Credential updated in database; session preserved.",
         "1. User opens Settings -> Profile.\n2. Inputs current password, new password, and confirms.\n3. Clicks Submit.\n4. Server validates current password and updates hash.", "scratch/seq_3.png"),
         
        ("UC104 - Sign Up", "Unauthorized User", "New user registers credentials.",
         "Anonymous user", "Account created and registered, redirects to Login page.",
         "1. User accesses sign up screen.\n2. Inputs email, name, password, and confirms.\n3. Submits form.\n4. Server saves registration details in database.", "scratch/seq_4.png"),
         
        ("UC105 - View Profile", "Authenticated User", "User views profile data details.",
         "User logged in", "Profile values loaded.",
         "1. User opens Profile tab.\n2. App fetches user data from DB using JWT.\n3. Displays profile sheet with full user details.", "scratch/seq_5.png"),
         
        ("UC106 - Edit Profile", "Authenticated User", "User edits contact details.",
         "User logged in", "Profile details updated.",
         "1. User clicks Edit Profile on details sheet.\n2. Updates email/phone number fields and clicks Save.\n3. Server commits changes to DB and returns updated JSON.", "scratch/seq_6.png"),
         
        ("UC201 - View / Search Student", "Admin / Teacher", "Search student registrations details.",
         "User authenticated", "Matching students list returned.",
         "1. User enters Search text in Student field.\n2. Backend queries database using search filter.\n3. Displays filtered list of matching students.", "scratch/seq_7.png"),
         
        ("UC202 - View / Search Teacher", "Admin", "Search teacher details.",
         "Admin authenticated", "Matching teachers list returned.",
         "1. Admin enters Search text in Teacher module.\n2. Backend queries database using search filter.\n3. Displays filtered list of matching teachers.", "scratch/seq_8.png"),
         
        ("UC203 - Add Class", "Admin", "Create a new classroom mapping.",
         "Admin authenticated", "Class table row inserted.",
         "1. Admin accesses Class Config module.\n2. Enters class Name, Grade, and selects Homeroom Teacher.\n3. Clicks Save. Server inserts Class row in database.", "scratch/seq_9.png"),
         
        ("UC204 - Edit Class", "Admin", "Modify class configuration.",
         "Admin authenticated", "Class row updated.",
         "1. Admin selects a class and clicks Edit.\n2. Modifies classroom, Homeroom teacher, or Grade.\n3. Clicks Save. Server updates database row.", "scratch/seq_10.png"),
         
        ("UC205 - View Timetable / Marks", "Student / Parent", "Check timetable schedules and subject scores.",
         "User authenticated", "Timetable grid and mark cards populated.",
         "1. User selects Semester and Week dropdown options.\n2. Renders weekly timetable schedule and subject exam scores fetched from DB.", "scratch/seq_11.png"),
         
        ("UC206 - Record Attendance / Marks", "Teacher", "Mark attendance or input grades.",
         "Teacher authenticated", "StudentMarks or AttendanceRecords row updated.",
         "1. Teacher selects Teaching class, scheduled date, and session.\n2. Inputs presence status or decimal marks score.\n3. Clicks Save. Server commits transaction.", "scratch/seq_12.png")
    ]

    for name, actor, desc, pre, post, flow, seq_img in uc_specs:
        add_h2(name)
        t_spec = doc.add_table(rows=6, cols=2)
        t_spec.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_spec)
        spec_rows = [
            ("Use Case ID & Name:", name),
            ("Primary Actor:", actor),
            ("Description:", desc),
            ("Pre-condition:", pre),
            ("Post-condition:", post),
            ("Normal Flow:", flow)
        ]
        for row_idx, (lbl, text) in enumerate(spec_rows):
            cell_lbl = t_spec.cell(row_idx, 0)
            cell_txt = t_spec.cell(row_idx, 1)
            set_cell_margins(cell_lbl, top=60, bottom=60, left=80, right=80)
            set_cell_margins(cell_txt, top=60, bottom=60, left=80, right=80)
            cell_lbl.paragraphs[0].add_run(lbl).bold = True
            cell_txt.paragraphs[0].add_run(text)
            
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(6)
        
        # Insert Sequence Diagram
        add_body(f"Sequence Flow Diagram for {name}:", italic=True)
        if os.path.exists(seq_img):
            doc.add_picture(seq_img, width=Inches(5.0))
            p_cap_seq = doc.add_paragraph()
            p_cap_seq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap_seq.add_run(f"Figure: {name} Sequence Diagram").italic = True
        else:
            add_body("[Sequence chart image missing]", italic=True)
            
        doc.add_page_break()

    # ─── SECTION 3: TABLE DEFINITIONS (All 35 Tables dynamically queried!) ───
    add_h1("3. Table Definitions (Database Schema)")
    add_body("The database holds 35 tables. The precise schemas, catalog data types, nullability, default values, and foreign keys are listed below:")

    for t_name, cols in db_tables.items():
        add_h2(f"Table: {t_name}")
        table_pks = pks.get(t_name, set())
        table_fks = fks.get(t_name, {})
        
        t_w = doc.add_table(rows=len(cols)+1, cols=5)
        t_w.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_w)
        headers_sc = ["Field Name", "Data Type", "Null?", "Default", "Key / Reference Table"]
        for col_idx, text in enumerate(headers_sc):
            cell = t_w.cell(0, col_idx)
            set_cell_background(cell, "1E3A8A")
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
            
        for row_idx, col in enumerate(cols):
            c_name = col.get("COLUMN_NAME")
            d_type = col.get("DATA_TYPE")
            c_len = col.get("CHARACTER_MAXIMUM_LENGTH")
            is_null = col.get("IS_NULLABLE")
            d_val = col.get("COLUMN_DEFAULT")
            
            formatted_type = f"{d_type}({c_len})" if c_len and int(c_len) > 0 else d_type
            if c_len == "-1":
                formatted_type = f"{d_type}(max)"
                
            key_status = ""
            if c_name in table_pks:
                key_status = "PK"
            if c_name in table_fks:
                ref_t = table_fks[c_name]
                key_status = f"FK (References {ref_t})"
                
            row_data = [c_name, formatted_type, is_null, d_val, key_status]
            for col_idx, text in enumerate(row_data):
                cell = t_w.cell(row_idx + 1, col_idx)
                set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
                p = cell.paragraphs[0]
                r = p.add_run(text)
                r.font.size = Pt(8.5)
                r.font.color.rgb = DARK_TEXT
                if col_idx == 0:
                    r.bold = True
                    
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(6)

    doc.add_page_break()

    # ─── SECTION 4: ERD ───
    add_h1("4. ERD (Entity Relationship Diagram)")
    add_body("Below is the visual Entity Relationship Diagram mapping logical entities and their cardinality mappings:")
    doc.add_picture("scratch/erd_diagram.png", width=Inches(6.0))
    p_cap_erd = doc.add_paragraph()
    p_cap_erd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_erd.add_run("Figure 4.1: Logical Entity Relationship Diagram").italic = True

    add_h2("Task Sheet Review 2")
    t_wbs2 = doc.add_table(rows=4, cols=4)
    t_wbs2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_wbs2)
    for col_idx, text in enumerate(wbs1_headers):
        cell = t_wbs2.cell(0, col_idx)
        set_cell_background(cell, "0A2540")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    wbs2_rows = [
        ("Đỗ Nguyễn Thiện Hoàng", "Architecture Design & Level 0 DFD Diagram", "Week 2", "Completed"),
        ("Trần Quang Khải", "Use Case Specifications & Sequence Diagrams", "Week 2", "Completed"),
        ("Phan Văn Duy", "Database Schema, Indexes & Entity Relationship Diagram (ERD)", "Week 2", "Completed")
    ]
    for row_idx, row_info in enumerate(wbs2_rows):
        for col_idx, text in enumerate(row_info):
            cell = t_wbs2.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].add_run(text)

    doc.add_page_break()

    # ─── REVIEW 3 SECTION ───
    add_review_header("REVIEW 3")

    # ─── SECTION 1: DATABASE DIAGRAM ───
    add_h1("1. Database Diagram")
    add_body("Below is the SQL Server relational schema mapping physical tables, columns, and foreign key connections:")
    doc.add_picture("scratch/database_diagram.png", width=Inches(6.0))
    p_cap_db = doc.add_paragraph()
    p_cap_db.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_db.add_run("Figure 5.2: Relational Database Schema Diagram").italic = True

    # ─── SECTION 2: WEB GUI DESIGN ───
    add_h1("2. Web GUI Design")
    add_body("Below are the web UI screenshots generated from the eStudiez application along with element mapping tables:")

    gui_mappings = [
        ("2.1 Home Landing Page", "media__1782016898396.png",
         "The eStudiez landing page displaying welcome bulletins and quick navigation shortcuts.",
         [("1", "News Bulletins Widget", "View updates", "Displays recent school announcements."),
          ("2", "Quick Links", "Click link", "Redirects to Timetable, Marks, and Chat portals."),
          ("3", "Helpdesk Button", "Click button", "Opens support message page.")]),
          
        ("2.2 Student Timetable", "media__1782036432724.png",
         "Displays the weekly timetable grid with period slots and subjects.",
         [("1", "Semester Dropdown", "Select option", "Filters slots by active semester."),
          ("2", "Timetable Grid Cell", "View details", "Shows subject, period time, and teacher name."),
          ("3", "Not Marked Status", "View badge", "Displays periods that have not yet had attendance recorded.")]),
          
        ("2.3 Student Attendance & Marks Statistics", "media__1782017160207.png",
         "Shows cumulative statistics cards and status color badges.",
         [("1", "Stats Summary Cards", "View numbers", "Shows counts of Present, Absent, Late, Excused."),
          ("2", "Attendance Grid", "View status", "Renders attendance status for each period across weeks."),
          ("3", "Semester Tab Selector", "Click tab", "Toggles statistics reports between Semester 1 and 2.")]),
          
        ("2.4 Student Marks / Progress Cards", "media__1782039833897.png",
         "Displays individual exam scores, weights, and detailed teacher evaluations.",
         [("1", "Subject Card", "Click header", "Expands to show detailed exam grading items."),
          ("2", "Score value", "View grade", "Shows numeric mark (0.0 to 10.0) and weights."),
          ("3", "Teacher Remarks", "View text", "Shows AI evaluation and study recommendations.")]),
          
        ("2.5 Study Resources Page", "media__1782117106269.png",
         "Repository of uploaded learning materials and homework documents.",
         [("1", "Search field", "Input text", "Filters documents by subject name or keywords."),
          ("2", "Download Button", "Click button", "Triggers browser download of target PDF or DOCX file."),
          ("3", "Resource Card", "Click link", "Opens document resource description panel.")]),
          
        ("2.6 User Registration (Sign Up)", "media__1782036432724.png",
         "Admin portal to register new students and issue codes.",
         [("1", "Name input box", "Input text", "Enter student's full name."),
          ("2", "Date of Birth", "Select date", "Standard date picker calendar control."),
          ("3", "Register Button", "Click button", "Validates input and inserts row in Student database.")]),
          
        ("2.7 User Login Portal", "media__1782017160207.png",
         "Authenticates students, teachers, and parents via JWT tokens.",
         [("1", "Username input box", "Input text", "Enter registered username or school email."),
          ("2", "Password input box", "Input text", "Enter encrypted password string."),
          ("3", "Sign In Button", "Click button", "Submits request to secure auth controller.")]),
          
        ("2.8 Admin Dashboard & Analytics", "media__1782039211975.png",
         "Displays system statistics, database summary, and news posts controls.",
         [("1", "User Statistics counters", "View dashboard", "Shows active students, teachers, classes."),
          ("2", "Petitions Grid", "Click row", "Admin views parent/student feedback ticket details."),
          ("3", "Create News Button", "Click button", "Opens rich text editor to write bulletin posts.")]),
          
        ("2.9 Admin User Accounts Management (CRUD Students)", "media__1782030886770.png",
         "CRUD list panel to manage student records.",
         [("1", "Search user bar", "Input text", "Filters accounts by code, name, or role."),
          ("2", "Add Student Button", "Click button", "Opens registration modal dialog."),
          ("3", "Action Menu Dropdown", "Click option", "Offers Edit Details, Reset Password, or Lock Account.")]),
          
        ("2.10 Admin Class Configuration Panel", "media__1782030893768.png",
         "Configuration page for high school grade cohorts.",
         [("1", "Cohort List grid", "View list", "Lists classes (e.g. 10A1, 10A2) configured for the school year."),
          ("2", "Assign Homeroom Teacher", "Select option", "Maps a teacher UUID to HomeroomTeacherId."),
          ("3", "Save Class Configuration", "Click button", "Updates the database Classes table.")]),
          
        ("2.11 Admin Teacher-Class Assignment Mapper", "media__1782035119584.png",
         "Maps subject teachers to specific timetable slots and classes.",
         [("1", "Teacher selector dropdown", "Click option", "Select from active teachers list."),
          ("2", "Subject-Class Mapper", "Select options", "Select Grade cohort and subject name to associate."),
          ("3", "Assign Button", "Click button", "Commits assignment to TeacherClassAssignments table.")]),
          
        ("2.12 Admin News Rich Publisher", "media__1782117106269.png",
         "Rich-text editor interface for school bulletins.",
         [("1", "Title Input field", "Input text", "Enter bulletin heading text."),
          ("2", "Rich-Text Textarea", "Input text", "Write announcements with rich formatting."),
          ("3", "Publish Button", "Click button", "Saves post to NewsPosts table, sending alert signals.")]),
          
        ("2.13 Teacher Period Attendance Sheet", "media__1782032322898.png",
         "Interface for teachers to mark presence status and notes.",
         [("1", "Roster Student Grid", "View list", "Displays student names and codes for the class."),
          ("2", "Attendance Status radios", "Select option", "Choose Present, Late, Absent, Excused for each student."),
          ("3", "Teacher note input", "Input text", "Add custom remarks about student behaviors or status notes.")]),
          
        ("2.14 Teacher Grading & AI Evaluation", "media__1782033130240.png",
         "Grading sheet equipped with AI Auto-Fill evaluation helper.",
         [("1", "Roster list panel", "View roster", "Displays student marks entry fields."),
          ("2", "AI Auto-Fill Button", "Click button", "Queries Spring Boot AI service for study directions."),
          ("3", "Remark Preview box", "View output", "Renders AI recommendations before saving.")])
    ]

    for title, filename, desc_txt, items in gui_mappings:
        add_h2(title)
        add_body(desc_txt)
        img_path = os.path.join(art_dir, filename)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
            p_cap_gui = doc.add_paragraph()
            p_cap_gui.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap_gui.add_run(f"Figure: Mockup of {title}").italic = True
        else:
            add_body("[Screenshot file missing or renamed]", italic=True)
            
        t_screen = doc.add_table(rows=len(items)+1, cols=4)
        t_screen.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_screen)
        headers_sc = ["No", "UI Element Name", "User Event", "Constraint / Description"]
        for col_idx, text in enumerate(headers_sc):
            cell = t_screen.cell(0, col_idx)
            set_cell_background(cell, "1E3A8A")
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)
            
        for row_idx, row_info in enumerate(items):
            for col_idx, text in enumerate(row_info):
                cell = t_screen.cell(row_idx + 1, col_idx)
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                p = cell.paragraphs[0]
                r = p.add_run(text)
                r.font.size = Pt(9)
                r.font.color.rgb = DARK_TEXT
                if col_idx == 0 or col_idx == 1: r.bold = True
                    
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(12)

    doc.add_page_break()

    # ─── SECTION 3: MOBILE GUI DESIGN ───
    add_h1("3. Mobile GUI Design")
    add_body("Below are mobile UI screen elements maps mapping events and descriptions:")

    mobile_mappings = [
        ("3.1 Mobile Login Screen", "mobile_login_v2_1782284921557.png",
         "The Login screen serves as the secure entry gateway for students, parents, and teachers accessing the eStudiez portal on mobile. It intercepts unauthenticated requests and requests verification via JWT. There is no public registration button as accounts are strictly pre-allocated by the school administration for safety.",
         [("1", "Logo Brand Icon", "View Item", "Visible at the top center of the screen.", "Displays the eStudiez purple gradient lightning-bolt logo to ensure brand recognition and app security."),
          ("2", "Username Input Field", "Input Text", "Must match the school-assigned username/email format (e.g., student.name@estudiez.edu.vn).", "The user enters their registered credential username or school-issued email to initiate authentication."),
          ("3", "Password Input Field", "Input Text", "Masked character display. Length between 6 to 32 characters.", "The user enters their account password corresponding to their username."),
          ("4", "Sign In Button", "Click Button", "Enabled only when username and password fields are not empty.", "Submits the credentials to the backend Spring Boot REST API (/api/auth/login) which responds with a JWT token upon successful authentication."),
          ("5", "Administrator Notice Label", "View Text", "Static read-only text at the bottom footer.", "Informs users that account creation is managed by school administrators to prevent unauthorized student registrations.")]),
          
        ("3.2 Mobile Dashboard Home", "mobile_dashboard_v2_1782284936073.png",
         "The Mobile Dashboard acts as the central hub for students, parents, and teachers, providing a summary of the current day's schedule, key notifications, and school announcements for rapid navigation.",
         [("1", "Header Brand Logo & Label", "View Item", "Fixed in the top navigation header.", "Renders the purple-blue lightning logo and active user profile avatar."),
          ("2", "Today's Timetable Widget", "Click Card", "Updates dynamically based on the current system date.", "Displays the student's scheduled classes and periods for today. Clicking it navigates to the full weekly timetable grid."),
          ("3", "Announcement Carousel Slider", "Swipe Banner / Click Slide", "Refreshes every 24 hours or on pull-to-refresh.", "Cycles through recent news headlines and events published by school administrators. Clicking a banner opens the rich news article detail."),
          ("4", "Home Tab Icon", "Click Icon", "Highlighted when active.", "Returns the user to this main dashboard home screen."),
          ("5", "Timetable Tab Icon", "Click Icon", "Provides access to student schedule.", "Redirects the user to the full academic class scheduling view."),
          ("6", "Notifications Tab Icon", "Click Icon", "Displays a badge count for unread notifications.", "Opens the notification feed for announcements, assignment updates, or attendance alerts."),
          ("7", "Profile Tab Icon", "Click Icon", "Provides access to account configurations.", "Redirects the user to the profile summary and settings panel.")]),
          
        ("3.3 Mobile Student Progress", "mobile_progress_v2_1782284952189.png",
         "The Student Progress screen provides real-time transparency for parents and students regarding academic grades and attendance records, encouraging proactive learning and monitoring.",
         [("1", "GPA Statistics Summary Card", "View Panel", "Calculates average grade point based on active semester marks.", "Renders the student's cumulative Grade Point Average (GPA) in a circular progress indicator."),
          ("2", "Attendance Ratio Card", "View Panel", "Derived from total period logs (Present / Total Classes).", "Displays the attendance percentage (e.g. 95%) to monitor overall student class attendance."),
          ("3", "Semester Selection Dropdown", "Select Dropdown Item", "Defaults to the current active school semester.", "Toggles the academic statistics and course lists between Semester 1 and Semester 2."),
          ("4", "Subject Grade List Card", "Click List Header", "Scrollable list view.", "Lists enrolled subjects (e.g. Mathematics, Chemistry, English) alongside their current grade letter and numerical progress bar. Clicking a subject reveals assignment and exam breakdown details.")]),
          
        ("3.4 Mobile User Profile", "mobile_profile_v2_1782284965957.png",
         "The User Profile screen houses account identity details, security change tools, and the exit interface to clear active JWT sessions.",
         [("1", "User Identity Header", "View Item", "Renders student profile picture from server database.", "Displays the student's full name, unique student ID, and registered academic class."),
          ("2", "Information Cards list", "View Text", "Read-only text.", "Lists account attributes including registered Email, Parent contact number, and Date of birth."),
          ("3", "Change Password Button", "Click Link", "Requires verification of current password.", "Directs the user to a secure modal interface to update their password."),
          ("4", "Sign Out Button", "Click Button", "Clears local JWT cache upon execution.", "Terminates the mobile session, purges the stored JWT token, and redirects the user back to the login screen.")])
    ]

    for m_title, m_file, m_desc, items in mobile_mappings:
        add_h2(m_title)
        add_body(m_desc)
        img_path = os.path.join(art_dir, m_file)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(3.2))
            p_cap_gui = doc.add_paragraph()
            p_cap_gui.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap_gui.add_run(f"Figure: Mobile Mockup of {m_title}").italic = True
        else:
            add_body("[Mobile screenshot file missing]", italic=True)
            
        t_mob = doc.add_table(rows=len(items)+1, cols=5)
        t_mob.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_mob)
        headers_sc_mob = ["No", "Name", "Event", "Constraint", "Description"]
        for col_idx, text in enumerate(headers_sc_mob):
            cell = t_mob.cell(0, col_idx)
            set_cell_background(cell, "1E3A8A")
            cell.paragraphs[0].add_run(text).bold = True
        for row_idx, row_info in enumerate(items):
            for col_idx, text in enumerate(row_info):
                cell = t_mob.cell(row_idx + 1, col_idx)
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                p = cell.paragraphs[0]
                r = p.add_run(text)
                r.font.size = Pt(9)
                r.font.color.rgb = DARK_TEXT
                if col_idx == 0 or col_idx == 1: r.bold = True
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(12)

    doc.add_page_break()

    # ─── SECTION 4: TASK SHEET REVIEW 3 ───
    add_h2("Task Sheet Review 3")
    t_wbs3 = doc.add_table(rows=4, cols=4)
    t_wbs3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_wbs3)
    for col_idx, text in enumerate(wbs1_headers):
        cell = t_wbs3.cell(0, col_idx)
        set_cell_background(cell, "0A2540")
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    wbs3_rows = [
        ("Đỗ Nguyễn Thiện Hoàng", "Spring Boot Backend REST APIs, JPA mapping & Integration Testing", "Weeks 3-5", "Completed"),
        ("Trần Quang Khải", "Vite Frontend React Screens, Routing & UI Integration", "Weeks 3-4", "Completed"),
        ("Phan Văn Duy", "Database Seeding SQL Scripts, Index Testing & Deployment", "Weeks 3-5", "Completed")
    ]
    for row_idx, row_info in enumerate(wbs3_rows):
        for col_idx, text in enumerate(row_info):
            cell = t_wbs3.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].add_run(text)

    # ─── PROJECT TASK SHEET ───
    add_h2("Project Task Sheet")
    add_body("Below is the daily task schedule for the entire development timeline:")
    t_pts = doc.add_table(rows=24, cols=5)
    t_pts.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_pts)
    pts_headers = ["Member", "Description Task", "Start Date", "End Date", "Status"]
    for col_idx, text in enumerate(pts_headers):
        cell = t_pts.cell(0, col_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    pts_rows = [
        ("Hoàng", "Teacher Backend (Springboot) Setup", "14/06/2026", "18/06/2026", "Complete"),
        ("Hoàng", "Attendance record JPA controller", "18/06/2026", "20/06/2026", "Complete"),
        ("Hoàng", "Marks grading database controller", "20/06/2026", "23/06/2026", "Complete"),
        ("Hoàng", "Study resources JPA uploads API", "23/06/2026", "25/06/2026", "Complete"),
        ("Hoàng", "Class Notifications broadcast system", "25/06/2026", "27/06/2026", "Complete"),
        ("Hoàng", "Teacher Frontend (Vite React) Setup", "01/07/2026", "06/07/2026", "Complete"),
        ("Hoàng", "Attendance Grid UI view and radios", "06/07/2026", "10/07/2026", "Complete"),
        ("Hoàng", "Marks grading roster and accordion", "10/07/2026", "15/07/2026", "Complete"),
        ("Hoàng", "AI learning path Auto-Fill button", "15/07/2026", "17/07/2026", "Complete"),
        ("Hoàng", "Mobile App (Flutter) teacher views", "17/07/2026", "31/07/2026", "Complete"),
        ("Khải", "Admin Backend (Springboot) Setup", "14/06/2026", "18/06/2026", "Complete"),
        ("Khải", "User Management CRUD JPA mapping", "18/06/2026", "22/06/2026", "Complete"),
        ("Khải", "Class Configuration & Assignings API", "22/06/2026", "25/06/2026", "Complete"),
        ("Khải", "News bulletins & events publishing", "25/06/2026", "27/06/2026", "Complete"),
        ("Khải", "Admin Frontend (Vite React) Setup", "01/07/2026", "05/07/2026", "Complete"),
        ("Khải", "User accounts list & edit modals", "05/07/2026", "10/07/2026", "Complete"),
        ("Khải", "Class assignment mappings selectors", "10/07/2026", "15/07/2026", "Complete"),
        ("Khải", "News posts rich editor integration", "15/07/2026", "18/07/2026", "Complete"),
        ("Khải", "Mobile App (Flutter) admin dashboards", "18/07/2026", "31/07/2026", "Complete"),
        ("Duy", "Student Backend (Springboot) Setup", "14/06/2026", "18/06/2026", "Complete"),
        ("Duy", "Timetable fetch endpoints", "18/06/2026", "22/06/2026", "Complete"),
        ("Duy", "Marks & attendance results API", "22/06/2026", "26/06/2026", "Complete"),
        ("Duy", "Resources download controllers", "26/06/2026", "30/06/2026", "Complete")
    ]
    for row_idx, row_info in enumerate(pts_rows):
        for col_idx, text in enumerate(row_info):
            cell = t_pts.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = DARK_TEXT

    doc.save("eStudiez_Design_Document.docx")
    print("Complete structured Word document successfully generated!")

if __name__ == "__main__":
    main()
