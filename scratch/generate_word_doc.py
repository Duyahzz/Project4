import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_element(name):
    return OxmlElement(name)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette
    PRIMARY_COLOR = RGBColor(10, 37, 64)   # Deep Blue
    SECONDARY_COLOR = RGBColor(30, 58, 138) # Navy Blue
    DARK_TEXT = RGBColor(51, 51, 51)       # Dark Grey

    # Cover Page
    # Center & Address Header
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_hdr1 = p_hdr.add_run("FPT ACADEMY INTERNATIONAL\nFPT – APTECH COMPUTER EDUCATION\n")
    r_hdr1.bold = True
    r_hdr1.font.size = Pt(13)
    r_hdr1.font.color.rgb = PRIMARY_COLOR
    r_hdr1.font.name = "Arial"
    
    r_hdr2 = p_hdr.add_run("Centre Name: ACE-HCMC-2-FPT\nAddress: 391A Nam Ky Khoi Nghia, District 3, Ho Chi Minh City, Viet Nam\n\n\n\n\n")
    r_hdr2.font.size = Pt(10)
    r_hdr2.italic = True
    r_hdr2.font.color.rgb = DARK_TEXT
    r_hdr2.font.name = "Arial"

    # Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("eStudiez (eStudent)\n")
    r_title.bold = True
    r_title.font.size = Pt(32)
    r_title.font.color.rgb = PRIMARY_COLOR
    r_title.font.name = "Arial"
    
    r_sub = p_title.add_run("(Next-Gen High School Student & Portal Management System)\n\n")
    r_sub.bold = True
    r_sub.font.size = Pt(18)
    r_sub.font.color.rgb = SECONDARY_COLOR
    r_sub.font.name = "Arial"

    r_doc = p_title.add_run("Project Design & Review Document\n\n\n\n\n")
    r_doc.italic = True
    r_doc.font.size = Pt(14)
    r_doc.font.color.rgb = DARK_TEXT
    r_doc.font.name = "Arial"

    # Group Info Block
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_info = [
        ("Supervisor:", "TRẦN PHƯỚC SINH"),
        ("Semester:", "IV"),
        ("Batch No:", "T1.2208,M2"),
        ("Group No:", "01"),
        ("Month:", "06 Year: 2026")
    ]
    for label, val in table_info:
        r_lbl = p_info.add_run(f"{label}  ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_lbl.font.color.rgb = PRIMARY_COLOR
        r_lbl.font.name = "Arial"
        
        r_val = p_info.add_run(f"{val}\n")
        r_val.font.size = Pt(11)
        r_val.font.color.rgb = DARK_TEXT
        r_val.font.name = "Arial"

    p_info.add_run("\n\n")

    # Group Members Table on Cover Page
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    headers = ["Order", "Full Name", "Roll No."]
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_background(cell, "0A2540")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        r.font.name = "Arial"

    members = [
        ("1.", "Hua Truong An", "Student1414334"),
        ("2.", "Le Ba Thanh", "Student1409136"),
        ("3.", "Nguyen Trung Quan", "Student1414325"),
        ("4.", "Bui Tran Anh Tri", "Student1414230")
    ]
    for row_idx, data in enumerate(members):
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT
            r.font.name = "Arial"

    doc.add_page_break()

    # Helpers for headings and text
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = PRIMARY_COLOR
        r.font.name = "Arial"
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = SECONDARY_COLOR
        r.font.name = "Arial"
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = SECONDARY_COLOR
        r.font.name = "Arial"
        return p

    def add_body(text, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_TEXT
        r.font.name = "Arial"
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_TEXT
        r.font.name = "Arial"
        return p

    # ─── ACKNOWLEDGEMENTS ───
    add_h1("Acknowledgements")
    add_body("The satisfaction that accompanies the successful completion of any task would be incomplete without the mention of people whose ceaseless cooperation made it possible, whose constant guidance and encouragement crown all efforts with success.")
    add_body("We are immensely grateful to our project guide, Mr. Tran Phuoc Sinh, for his guidance, inspiration, and constructive suggestions that significantly contributed to the preparation and successful completion of our study-progress tracking project. His expertise and support were invaluable throughout the process.")
    add_body("We also extend our heartfelt thanks to our colleagues who contributed their time, effort, and insights, which were crucial to the successful completion of this project. Their collaboration and dedication played a pivotal role in overcoming challenges and achieving our goals.")
    doc.add_page_break()

    # ─── SECTION 1: OVERVIEW ───
    add_h1("1. Overview")
    
    add_h2("1.1 Project Information")
    add_bullet("Project name: eStudiez (eStudent) — A Study-Progress Tracking Application")
    add_bullet("Group name: Group 01")
    add_bullet("Software type: Web Application")
    
    add_h2("1.2 Product Background")
    add_body("Designing and developing a School Study-Progress Tracking Web Application involves several stages of production. Here's a brief overview of the production background for such a project:")
    add_body("Requirement Gathering: This stage involves identifying the requirements of both the school board (Admins), subject teachers, students, and parents/guardians for the application. It is essential to gather feedback from users to determine the features and functionality that should be included in the application.")
    add_body("Design: Once the requirements are clear, the design phase begins. It includes creating the visual design of the application and defining the user interface and user experience. A prototype may be developed to get feedback from users.")
    add_body("Development: This phase involves the actual development of the application. The development team uses the appropriate technology stack to build the application, and the features identified in the design stage will be implemented.")
    add_body("Testing: Before the application is released, it must be thoroughly tested to ensure that it meets the requirements and functions as expected. Various types of testing, including functional, performance, and security testing, should be conducted.")
    
    add_h2("1.3 Existing Systems")
    add_body("There are several existing systems for study progress and school portal management that we considered when developing this app:")
    add_body("vnEdu: An electronic school report system in Vietnam. It provides portals for teachers to input marks and parents to view them, but lacks integrated real-time class chat rooms and AI-driven study path generators.")
    add_body("SMAS: A robust school management software developed by Viettel. It is heavy on administrative processes and school statistics but is less streamlined for interactive student-teacher communication and sharing study resources.")
    add_body("Canvas LMS: A widely adopted learning management system. While excellent for course content, it is built for higher education and is not configured natively for Vietnamese high school class structures and homeroom/parent links.")
    
    add_h2("1.4 Business & Educational Opportunity")
    add_body("E-commerce has revolutionized business operations, and similarly, portal centralization has modernized school operations. Key opportunities include:")
    add_bullet("Centralized Academic Portal: Providing a single platform for tracking subject marks, attendance, and timetables.")
    add_bullet("Parent Engagement: Connecting parents directly with classrooms via Parent-Teacher chat groups and child progress notifications.")
    add_bullet("Resource Accessibility: Online subject resources (documents, videos, links) ensure students can study remotely 24/7.")
    add_bullet("AI Academic Insights: Generating automated learning recommendations for students based on their exam performance and weaknesses.")

    add_h2("1.5 Project Scope & Limitations")
    add_body("Scope: Academic progress tracking (marks & attendance), timetables, news bulletins, notifications, class chat groups, extra revision classes, study resources repository, and parent contact management.")
    add_body("Limitations: No teacher HR management (payroll/recruitment), no admissions processing, no student fee transactions, and restricted to high school class structures (grades 10, 11, and 12).")
    doc.add_page_break()

    # ─── SECTION 2: USER REQUIREMENTS ───
    add_h1("2. User Requirements")
    
    add_h2("2.1 Actor Descriptions")
    table_act = doc.add_table(rows=5, cols=2)
    table_act.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_act)
    
    headers_act = ["Actor", "Description"]
    for col_idx, text in enumerate(headers_act):
        cell = table_act.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        
    actors = [
        ("Admin", "Super Admin (BGH). Manages user creation, homeroom classes, subject-teacher assignments, and posts school news."),
        ("Teacher", "Subject teachers. Record attendance per period, manage student marks, upload resources, and write AI evaluations."),
        ("Student", "Students of grades 10/11/12. View timetables, track personal marks/attendance, download materials, and chat with teachers."),
        ("Parent", "Guardians linked to students. Monitor child's marks and attendance, view school bulletins, and join Parent-Teacher chats.")
    ]
    for row_idx, data in enumerate(actors):
        for col_idx, text in enumerate(data):
            cell = table_act.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT

    add_h2("2.2 Data Flow Diagram (DFD) Level 0")
    add_body("The Level 0 DFD illustrates the flow of information between actors and the central system:")
    add_bullet("Admin inputs: Student/Teacher accounts, Class details, School News.")
    add_bullet("Teacher inputs: Attendance statuses, Exam marks, Resource files, AI prompt metrics.")
    add_bullet("Student/Parent outputs: Timetable grid, Mark list, Notification alerts, Chat message feeds.")
    
    add_h2("2.3 Use Case List")
    table_uc = doc.add_table(rows=11, cols=3)
    table_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_uc)
    
    headers_uc = ["UC Code", "Use Case Name", "Description"]
    for col_idx, text in enumerate(headers_uc):
        cell = table_uc.cell(0, col_idx)
        set_cell_background(cell, "0B2548")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    usecases = [
        ("UC101", "Login / Logout", "Authenticates all users using encrypted passwords."),
        ("UC102", "Change Password", "Allows users to update their credentials for safety."),
        ("UC103", "View/Edit Profile", "View and update email, phone, and profile avatar."),
        ("UC201", "View Timetable", "Students/Parents view class weekly schedule grid."),
        ("UC202", "View Progress", "Students/Parents track marks and attendance status."),
        ("UC203", "Download Resource", "Students browse and download subject study materials."),
        ("UC301", "Record Attendance", "Teachers mark student presence, lateness, or excuses."),
        ("UC302", "Update Marks", "Teachers record exam scores and detailed remarks."),
        ("UC303", "AI Evaluation", "AI processes scores to suggest custom study steps."),
        ("UC401", "Manage Users & Classes", "Admins create students, classes, and assign teachers.")
    ]
    for row_idx, data in enumerate(usecases):
        for col_idx, text in enumerate(data):
            cell = table_uc.cell(row_idx + 1, col_idx)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT

    doc.add_page_break()

    # ─── SECTION 3: MANAGEMENT & SYSTEM REQUIREMENTS ───
    add_h1("3. Management Plan")
    
    add_h2("3.1 Scope and WBS Estimation")
    add_body("The project was scheduled over 5 weeks (July 10 to August 18). Implementation details:")
    add_bullet("Phase 1: Database Setup & JPA mapping (2 days)")
    add_bullet("Phase 2: Authentication Service & JWT filters (3 days)")
    add_bullet("Phase 3: Core Academic APIs (Teacher & Student controllers) (5 days)")
    add_bullet("Phase 4: Responsive Web Front-end with React Context integration (10 days)")
    add_bullet("Phase 5: System Integration & Compile checks (5 days)")
    
    add_h2("3.2 DevOps Workflow")
    add_body("Applying the DevOps model to management enabled the group to implement the product with coordination among members. The pipeline consisted of local Git commits, GitHub code integration, type check verification via npm, and database migrations on SQL Server.")
    
    add_h1("4. System Requirements")
    add_h2("4.1 Hardware Requirements")
    add_bullet("Dual-Core Processor speed 2.0Ghz or faster.")
    add_bullet("8 GB of RAM recommended for database and web servers.")
    add_bullet("Hard disk space: minimum 50GB.")
    add_h2("4.2 Software Requirements")
    add_bullet("Database: Microsoft SQL Server 2019+.")
    add_bullet("IDE: IntelliJ IDEA 2023+ / VS Code.")
    add_bullet("Build Tools: JDK 17, Gradle 8.x, Node.js v20.")
    add_h2("4.3 Technology Summary")
    add_bullet("Web App: React, TypeScript, React Router 7, Tailwind CSS 4.")
    add_bullet("Backend: Spring Boot 3.x, Hibernate JPA, JWT Security.")
    doc.add_page_break()

    # ─── SECTION 4: TABLE DEFINITIONS ───
    add_h1("5. Table Definitions (Database Schema)")
    add_body("The eStudentDB database is designed with structured relational schemas. The main tables are defined below:")

    schema_data = {
        "Users": [
            ("UserId", "uniqueidentifier", "NO", "newid()", "Primary Key. Unique identifier."),
            ("RoleId", "int", "NO", "NULL", "FK to Roles table."),
            ("Username", "nvarchar(80)", "NO", "NULL", "Unique username for login."),
            ("PasswordHash", "nvarchar(255)", "NO", "NULL", "BCrypt hashed password."),
            ("FullName", "nvarchar(150)", "NO", "NULL", "Full name of the user."),
            ("Email", "nvarchar(150)", "YES", "NULL", "Optional email address."),
            ("Phone", "nvarchar(30)", "YES", "NULL", "Optional phone number."),
            ("AvatarUrl", "nvarchar(500)", "YES", "NULL", "Profile image URL."),
            ("IsActive", "bit", "NO", "((1))", "Account status flag.")
        ],
        "Students": [
            ("StudentId", "uniqueidentifier", "NO", "newid()", "Primary Key. Unique identifier."),
            ("UserId", "uniqueidentifier", "NO", "NULL", "FK to Users table."),
            ("StudentCode", "nvarchar(50)", "NO", "NULL", "Unique student registration code."),
            ("DateOfBirth", "date", "YES", "NULL", "Student's date of birth."),
            ("Gender", "nvarchar(20)", "YES", "NULL", "Student's gender."),
            ("Address", "nvarchar(max)", "YES", "NULL", "Residential address."),
            ("AdmissionDate", "date", "NO", "getdate()", "Date enrolled in the school.")
        ],
        "Teachers": [
            ("TeacherId", "uniqueidentifier", "NO", "newid()", "Primary Key. Unique identifier."),
            ("UserId", "uniqueidentifier", "NO", "NULL", "FK to Users table."),
            ("EmployeeCode", "nvarchar(50)", "NO", "NULL", "Unique teacher registration code."),
            ("SubjectId", "int", "NO", "NULL", "FK to Subjects table."),
            ("Qualification", "nvarchar(150)", "YES", "NULL", "Professional qualification.")
        ],
        "Classes": [
            ("ClassId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("SchoolYearId", "int", "NO", "NULL", "FK to SchoolYears table."),
            ("GradeId", "int", "NO", "NULL", "FK to Grades table."),
            ("Name", "nvarchar(50)", "NO", "NULL", "Classroom label (e.g. 10A1)."),
            ("HomeroomTeacherId", "uniqueidentifier", "YES", "NULL", "FK to Teachers table."),
            ("TrainingProgram", "nvarchar(30)", "NO", "'REGULAR'", "REGULAR or REVISION program."),
            ("Room", "nvarchar(50)", "YES", "NULL", "Assigned school room.")
        ],
        "TimetableSlots": [
            ("TimetableSlotId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("ClassId", "int", "NO", "NULL", "FK to Classes table."),
            ("SubjectId", "int", "NO", "NULL", "FK to Subjects table."),
            ("TeacherId", "uniqueidentifier", "NO", "NULL", "FK to Teachers table."),
            ("SemesterId", "int", "NO", "NULL", "FK to Semesters table."),
            ("DayOfWeek", "tinyint", "NO", "NULL", "Day index (1 = Mon to 6 = Sat)."),
            ("PeriodNo", "tinyint", "NO", "NULL", "Period position (1 to 5).")
        ],
        "LessonSessions": [
            ("LessonSessionId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("TimetableSlotId", "int", "YES", "NULL", "FK to TimetableSlots."),
            ("ClassId", "int", "NO", "NULL", "FK to Classes table."),
            ("SubjectId", "int", "NO", "NULL", "FK to Subjects table."),
            ("TeacherId", "uniqueidentifier", "NO", "NULL", "FK to Teachers table."),
            ("SessionDate", "date", "NO", "NULL", "Specific scheduled date."),
            ("PeriodNo", "tinyint", "NO", "NULL", "Period number."),
            ("Status", "nvarchar(30)", "NO", "'SCHEDULED'", "SCHEDULED or COMPLETED.")
        ],
        "AttendanceRecords": [
            ("AttendanceId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("LessonSessionId", "int", "NO", "NULL", "FK to LessonSessions."),
            ("StudentId", "uniqueidentifier", "NO", "NULL", "FK to Students table."),
            ("Status", "nvarchar(30)", "NO", "NULL", "PRESENT, ABSENT, LATE, EXCUSED."),
            ("Note", "nvarchar(max)", "YES", "NULL", "Optional description note."),
            ("RecordedBy", "uniqueidentifier", "NO", "NULL", "FK to Users table.")
        ],
        "Assessments": [
            ("AssessmentId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("ClassId", "int", "NO", "NULL", "FK to Classes table."),
            ("SubjectId", "int", "NO", "NULL", "FK to Subjects table."),
            ("TeacherId", "uniqueidentifier", "NO", "NULL", "FK to Teachers table."),
            ("SemesterId", "int", "NO", "NULL", "FK to Semesters table."),
            ("Title", "nvarchar(255)", "NO", "NULL", "Assessment title (e.g. Midterm 1)."),
            ("MaxScore", "decimal", "NO", "((10))", "Maximum possible score."),
            ("Weight", "decimal", "NO", "((1))", "Weight factor in report cards.")
        ],
        "StudentMarks": [
            ("StudentMarkId", "int", "NO", "NULL", "Primary Key. Identity seed."),
            ("AssessmentId", "int", "NO", "NULL", "FK to Assessments table."),
            ("StudentId", "uniqueidentifier", "NO", "NULL", "FK to Students table."),
            ("Score", "decimal", "NO", "NULL", "Earned score value."),
            ("TeacherComment", "nvarchar(max)", "YES", "NULL", "General feedback comment."),
            ("Remark", "nvarchar(max)", "YES", "NULL", "AI evaluation text recommendations.")
        ]
    }

    for t_name, cols in schema_data.items():
        add_h2(f"Table: {t_name}")
        t = doc.add_table(rows=len(cols)+1, cols=5)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        
        headers_sc = ["Field Name", "Data Type", "Null?", "Default", "Description"]
        for col_idx, text in enumerate(headers_sc):
            cell = t.cell(0, col_idx)
            set_cell_background(cell, "1E3A8A")
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)
            
        for row_idx, col_info in enumerate(cols):
            for col_idx, text in enumerate(col_info):
                cell = t.cell(row_idx + 1, col_idx)
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                p = cell.paragraphs[0]
                r = p.add_run(text)
                r.font.size = Pt(9)
                r.font.color.rgb = DARK_TEXT
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(6)

    doc.add_page_break()

    # ─── SECTION 5: USE CASE FLOWS ───
    add_h1("6. Use Case specifications & Sequence flow")
    
    add_h2("6.1 UC101 - Student Portal Login Flow")
    add_body("Actors: Student / Parent / Teacher / Admin")
    add_body("Normal Flow Steps:")
    add_bullet("1. User accesses the eStudiez website homepage.")
    add_bullet("2. Clicks the Sign In button to open the credential forms.")
    add_bullet("3. Inputs registered Username/Email and password values.")
    add_bullet("4. Clicks Continue. Request is forwarded to Spring Boot `/api/auth/login` endpoint.")
    add_bullet("5. The AuthManager queries database credentials and verifies the password hash.")
    add_bullet("6. Returns JWT payload. Client browser caches it and opens the Dashboard.")

    add_h2("6.2 UC301 - Record Attendance Flow")
    add_body("Actors: Subject Teacher")
    add_body("Normal Flow Steps:")
    add_bullet("1. Teacher logs in and opens the Attendance tab on the sidebar.")
    add_bullet("2. Selects active Teaching Class, Subject, Period slot, and date.")
    add_bullet("3. Renders the list of enrolled students fetched via ClassEnrollments.")
    add_bullet("4. Marks students Present, Late, Absent, or Excused, adding comments if needed.")
    add_bullet("5. Clicks Save Changes. Request is uploaded to `/api/lessons/{id}/attendance`.")
    add_bullet("6. Server saves to AttendanceRecords and shows 'Attendance Recorded!' toast.")

    add_h2("6.3 UC302 - Update Marks & AI Recommendations Flow")
    add_body("Actors: Subject Teacher")
    add_body("Normal Flow Steps:")
    add_bullet("1. Teacher selects an Exam (Quiz, Midterm, Final) under the Marks & Evaluation tab.")
    add_bullet("2. Roster is loaded with inputs for numeric scores and review comments.")
    add_bullet("3. Teacher inputs student grade and clicks the 'AI Auto-Fill' button.")
    add_bullet("4. Spring Boot forwards student metrics to the AI integration service.")
    add_bullet("5. AI model returns custom study directions, which populate the evaluation form.")
    add_bullet("6. Teacher reviews output and submits. Server records data to StudentMarks table.")

    doc.add_page_break()

    # ─── SECTION 6: ERD & GRAPHICS ───
    add_h1("7. System ERD & Database Relationships")
    add_body("The relationship structure between the main entities is defined as follows:")
    add_bullet("Users acts as the base entity for Student, Teacher, and Parent (1:1 relationships).")
    add_bullet("Class contains multiple Students linked via the ClassEnrollments join table (1:N).")
    add_bullet("TimetableSlots maps a Class to a Subject, Teacher, Day, and Period slot (M:N mapping).")
    add_bullet("LessonSessions are spawned from TimetableSlots, serving as parent sessions for AttendanceRecords (1:N).")
    add_bullet("Assessments are classroom-specific tests that group StudentMarks details (1:N).")
    add_bullet("StudentParentLinks connect Student records to their primary Parent guardians (M:N mapping).")

    add_h1("8. Web GUI Design")
    add_body("The user interfaces are divided into distinct dashboard portals:")
    add_bullet("Student Portal: Simple layout displaying current weekly timetables (highlighting active periods) and latest subject mark bulletins.")
    add_bullet("Teacher Portal: Structured data sheet displaying rosters to check attendance and grading forms equipped with the AI generation modal.")
    add_bullet("Parent Portal: Overview showing children's current-week schedule and summaries of active marks/absences.")
    add_bullet("Admin Portal: Comprehensive control board to add new users, perform class transfers, and broadcast school-wide news posts.")

    # Save Word Doc
    doc.save("eStudiez_Design_Document.docx")
    print("Word document generated successfully at: c:/zzz/AP-Project4/eStudiez_Design_Document.docx")

if __name__ == "__main__":
    build_document()
